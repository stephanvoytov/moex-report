import os
import time
import requests
import pandas as pd
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# =========================================================
# НАСТРОЙКИ
# =========================================================

INDEXES = {
    "IMOEX": {
        "ticker": "IMOEX",
        "name": "Индекс Мосбиржи IMOEX",
    },
    "RTSI": {
        "ticker": "RTSI",
        "name": "Индекс Мосбиржи РТС (RTSI)",
    },
    "RGBITR": {
        "ticker": "RGBITR",
        "name": "Индекс Мосбиржи государственных облигаций (RGBITR)",
    },
    "RUCBTRNS": {
        "ticker": "RUCBTRNS",
        "name": "Индекс Мосбиржи корпоративных облигаций (RUCBTRNS)",
    },
}

BASE_URL = "https://iss.moex.com/iss/history/engines/stock/markets/index/securities"
CBR_URL = "https://www.cbr.ru/scripts/XML_dynamic.asp"
USD_CODE = "R01235"

RUS_WEEKDAYS = {
    0: "Понедельник",
    1: "Вторник",
    2: "Среда",
    3: "Четверг",
    4: "Пятница",
    5: "Суббота",
    6: "Воскресенье",
}

RUS_MONTHS_GEN = {
    1: "января", 2: "февраля", 3: "марта", 4: "апреля",
    5: "мая", 6: "июня", 7: "июля", 8: "августа",
    9: "сентября", 10: "октября", 11: "ноября", 12: "декабря",
}

START_DATE = ""
END_DATE = ""
LOAD_FROM = ""
LOAD_TO = ""


# =========================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# =========================================================

def format_number(value, digits=2):
    if pd.isna(value):
        return "-"
    return f"{value:,.{digits}f}".replace(",", " ").replace(".", ",")


def format_percent(value):
    if pd.isna(value):
        return "-"

    if abs(value) < 0.005:
        return "0,00%"

    sign = "+" if value > 0 else ""
    return f"{sign}{value:.2f}%".replace(".", ",")


def short_date(date_str):
    if date_str is None:
        return "-"
    dt = datetime.strptime(date_str, "%Y-%m-%d")
    weekdays_short = {0: "пн", 1: "вт", 2: "ср", 3: "чт", 4: "пт", 5: "сб", 6: "вс"}
    return f"{dt.strftime('%d.%m')} ({weekdays_short[dt.weekday()]})"


def full_date(date_str):
    dt = datetime.strptime(date_str, "%Y-%m-%d")
    weekday = RUS_WEEKDAYS[dt.weekday()]
    day = dt.day
    month = RUS_MONTHS_GEN[dt.month]
    year = dt.year
    return f"{weekday} {day} {month} {year}"


def convert_to_cbr_date(date_str):
    dt = datetime.strptime(date_str, "%Y-%m-%d")
    return dt.strftime("%d.%m.%Y")


# =========================================================
# ЗАГРУЗКА КУРСА USD/RUB
# =========================================================

def make_session():
    session = requests.Session()
    retries = Retry(
        total=5,
        backoff_factor=2,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET"],
    )
    adapter = HTTPAdapter(max_retries=retries)
    session.mount("https://", adapter)
    session.mount("http://", adapter)
    return session


def load_usd_rates():
    date_from = convert_to_cbr_date(LOAD_FROM)
    date_to = convert_to_cbr_date(LOAD_TO)

    params = {
        "date_req1": date_from,
        "date_req2": date_to,
        "VAL_NM_RQ": USD_CODE,
    }

    try:
        session = make_session()
        response = session.get(CBR_URL, params=params, timeout=60)
        response.encoding = "utf-8"
        root = ET.fromstring(response.text)
    except Exception as e:
        print(f"Ошибка загрузки курса USD/RUB: {e}")
        return {}

    rates = {}
    for record in root.findall("Record"):
        date_str = record.get("Date")
        value_elem = record.find("Value")
        if date_str and value_elem is not None:
            raw = value_elem.text.strip().replace(",", ".")
            try:
                rate = round(float(raw), 4)
                dt = datetime.strptime(date_str, "%d.%m.%Y")
                rates[dt.strftime("%Y-%m-%d")] = rate
            except (ValueError, TypeError):
                continue

    return rates


# =========================================================
# ЗАГРУЗКА ДАННЫХ ИНДЕКСОВ
# =========================================================

def load_data(ticker):
    url = f"{BASE_URL}/{ticker}.json"

    params = {
        "from": LOAD_FROM,
        "till": LOAD_TO,
        "iss.meta": "off",
    }

    try:
        session = make_session()
        response = session.get(url, params=params, timeout=60)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(f"Ошибка загрузки данных для {ticker}: {e}")
        return pd.DataFrame()

    try:
        data = response.json()
        columns = data["history"]["columns"]
        rows = data["history"]["data"]
        df = pd.DataFrame(rows, columns=columns)
    except (KeyError, ValueError) as e:
        print(f"Ошибка обработки ответа для {ticker}: {e}")
        return pd.DataFrame()

    numeric_columns = ["OPEN", "LOW", "HIGH", "CLOSE", "VALUE", "DURATION", "YIELD"]

    for col in numeric_columns:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    return df


# =========================================================
# ОБРАБОТКА ДАННЫХ
# =========================================================

def prepare_dataframe(df):
    if df.empty:
        return df

    df = df.copy()
    df["CHANGE"] = df["CLOSE"].pct_change() * 100
    return df


def weekly_change(df):
    if df.empty:
        return None

    first_close = df.iloc[0]["CLOSE"]
    last_close = df.iloc[-1]["CLOSE"]
    return ((last_close / first_close) - 1) * 100


def find_max(df):
    if df.empty:
        return None, None

    idx = df["HIGH"].idxmax()
    return df.loc[idx, "HIGH"], df.loc[idx, "TRADEDATE"]


def find_min(df):
    if df.empty:
        return None, None

    idx = df["LOW"].idxmin()
    return df.loc[idx, "LOW"], df.loc[idx, "TRADEDATE"]


# =========================================================
# WORD — СТИЛИ ТАБЛИЦ
# =========================================================

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, val in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), val.get("val", "single"))
        element.set(qn("w:sz"), val.get("sz", "4"))
        element.set(qn("w:color"), val.get("color", "000000"))
        element.set(qn("w:space"), val.get("space", "0"))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tblPr.append(borders)


# =========================================================
# WORD — ФОРМАТИРОВАНИЕ
# =========================================================

def set_font(run, size=10, bold=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Times New Roman"


def add_title(document, text):
    p = document.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    set_font(run, 14, True)


def add_heading(document, text):
    p = document.add_paragraph()
    run = p.add_run(text)
    set_font(run, 12, True)


def add_table(document, headers, rows):
    num_cols = len(headers)
    if num_cols == 0 or not rows:
        return

    table = document.add_table(rows=1, cols=num_cols)
    remove_table_borders(table)

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_border(hdr_cells[i],
            top={"val": "single", "sz": "6"},
            bottom={"val": "single", "sz": "6"},
            left={"val": "single", "sz": "6"},
            right={"val": "single", "sz": "6"},
        )
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        set_font(run, 9, True)

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_border(row_cells[i],
                top={"val": "single", "sz": "6"},
                bottom={"val": "single", "sz": "6"},
                left={"val": "single", "sz": "6"},
                right={"val": "single", "sz": "6"},
            )
            p = row_cells[i].paragraphs[0]
            run = p.add_run(str(value))
            set_font(run, 9)

    document.add_paragraph()


# =========================================================
# WORD — ДНЕВНАЯ СВОДКА
# =========================================================

def add_daily_summary(doc, display_data, usd_rates):
    add_heading(doc, "Дневная сводка")

    first_key = next(iter(INDEXES.keys()))
    df_display = display_data[first_key]
    min_rows = min(len(display_data[k]) for k in INDEXES)

    for i in range(min_rows):
        ref_date = df_display.iloc[i]["TRADEDATE"]

        p = doc.add_paragraph()
        run = p.add_run(full_date(ref_date))
        set_font(run, 11, True)

        for key in INDEXES:
            row = display_data[key].iloc[i]
            close = row["CLOSE"]
            change = row["CHANGE"]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f"{key}: ")
            set_font(run, 10, True)
            run = p.add_run(f"{format_number(close)} ({format_percent(change)})")
            set_font(run, 10)

        usd_rate = usd_rates.get(ref_date)
        if usd_rate is None:
            dt = datetime.strptime(ref_date, "%Y-%m-%d")
            for offset in range(1, 5):
                prev = (dt - timedelta(days=offset)).strftime("%Y-%m-%d")
                if prev in usd_rates:
                    usd_rate = usd_rates[prev]
                    break

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run("USD/RUB: ")
        set_font(run, 10, True)
        if usd_rate is not None:
            rate_str = f"{usd_rate:,.4f}".replace(",", " ").replace(".", ",")
            run = p.add_run(f"{rate_str} (курс ЦБ)")
        else:
            run = p.add_run("—")
        set_font(run, 10)

        doc.add_paragraph()


# =========================================================
# WORD — СВОДНАЯ ТАБЛИЦА
# =========================================================

def add_summary_table(doc, all_data, display_data):
    add_heading(doc, "Сводная таблица")

    summary_headers = ["Дата", "IMOEX", "RTSI", "RGBITR", "RUCBTRNS"]
    summary_rows = []

    min_rows = min(len(display_data[k]) for k in INDEXES)
    ref_key = next(iter(INDEXES.keys()))

    for i in range(min_rows):
        row = [short_date(display_data[ref_key].iloc[i]["TRADEDATE"])]
        for key in INDEXES:
            r = display_data[key].iloc[i]
            row.append(f"{format_number(r['CLOSE'])} ({format_percent(r['CHANGE'])})")
        summary_rows.append(row)

    # Итог недели
    summary_rows.append([
        "Итог недели",
        format_percent(weekly_change(display_data["IMOEX"])),
        format_percent(weekly_change(display_data["RTSI"])),
        format_percent(weekly_change(display_data["RGBITR"])),
        format_percent(weekly_change(display_data["RUCBTRNS"])),
    ])

    # Максимум
    values = {}
    for key in INDEXES:
        v, d = find_max(display_data[key])
        values[key] = (v, d)

    summary_rows.append([
        "Максимум",
        f"{format_number(values['IMOEX'][0])} ({short_date(values['IMOEX'][1])})",
        f"{format_number(values['RTSI'][0])} ({short_date(values['RTSI'][1])})",
        f"{format_number(values['RGBITR'][0])} ({short_date(values['RGBITR'][1])})",
        f"{format_number(values['RUCBTRNS'][0])} ({short_date(values['RUCBTRNS'][1])})",
    ])

    # Минимум
    values = {}
    for key in INDEXES:
        v, d = find_min(display_data[key])
        values[key] = (v, d)

    summary_rows.append([
        "Минимум",
        f"{format_number(values['IMOEX'][0])} ({short_date(values['IMOEX'][1])})",
        f"{format_number(values['RTSI'][0])} ({short_date(values['RTSI'][1])})",
        f"{format_number(values['RGBITR'][0])} ({short_date(values['RGBITR'][1])})",
        f"{format_number(values['RUCBTRNS'][0])} ({short_date(values['RUCBTRNS'][1])})",
    ])

    add_table(doc, summary_headers, summary_rows)


# =========================================================
# АВТО-ДАТЫ (ПРОШЛАЯ НЕДЕЛЯ)
# =========================================================

def get_previous_week_dates():
    today = datetime.now()
    days_since_monday = today.weekday()
    current_monday = today - timedelta(days=days_since_monday)
    prev_monday = current_monday - timedelta(days=7)
    prev_friday = prev_monday + timedelta(days=4)
    return prev_monday.strftime("%Y-%m-%d"), prev_friday.strftime("%Y-%m-%d")


# =========================================================
# ОСНОВНАЯ ЛОГИКА
# =========================================================

def main():
    all_data = {}

    for key, info in INDEXES.items():
        print(f"Загружаю {info['name']}...")
        df = load_data(info["ticker"])
        df = prepare_dataframe(df)
        all_data[key] = df

    print("Загружаю курс USD/RUB...")
    usd_rates = load_usd_rates()

    display_data = {}
    for key in INDEXES:
        df = all_data[key]
        if df.empty:
            display_data[key] = df
        else:
            display_data[key] = df[df["TRADEDATE"] >= START_DATE].reset_index(drop=True)

    print("Формирую отчёт...")
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    add_title(doc, "Отчёт по индексам MOEX")

    # Дневная сводка
    if all(len(display_data[k]) > 0 for k in INDEXES):
        add_daily_summary(doc, display_data, usd_rates)

    # ПОДРОБНЫЕ ТАБЛИЦЫ
    for key, info in INDEXES.items():
        if display_data[key].empty:
            print(f"Нет данных для индекса {info['name']}")
            continue

        df = display_data[key]
        add_heading(doc, info["name"])

        headers = ["Дата", "open", "high", "low", "close", "Изм. %"]

        if key == "IMOEX":
            headers.append("Объём, млрд ₽")
        elif key == "RTSI":
            headers.append("Объём, млн $")
        else:
            headers.extend(["Дюрация, дн.", "YIELD, %", "Объём, млрд ₽"])

        rows = []
        for _, row in df.iterrows():
            values = [
                short_date(row["TRADEDATE"]),
                format_number(row["OPEN"]),
                format_number(row["HIGH"]),
                format_number(row["LOW"]),
                format_number(row["CLOSE"]),
                format_percent(row["CHANGE"]),
            ]

            if key == "IMOEX":
                values.append(format_number(row["VALUE"] / 1_000_000_000))
            elif key == "RTSI":
                values.append(format_number(row["VALUE"] / 1_000_000))
            else:
                values.extend([
                    format_number(row["DURATION"], 0),
                    format_number(row["YIELD"]),
                    format_number(row["VALUE"] / 1_000_000_000),
                ])

            rows.append(values)

        add_table(doc, headers, rows)

    # СВОДНАЯ ТАБЛИЦА
    if all(len(display_data[k]) > 0 for k in INDEXES):
        add_summary_table(doc, all_data, display_data)

    # СОХРАНЕНИЕ
    now = datetime.now()
    filename = f"moex_report_{now:%Y%m%d_%H%M}.docx"

    try:
        doc.save(filename)
        print(f"Файл сохранён: {filename}")
        os.startfile(filename)
    except Exception as e:
        print(f"Ошибка сохранения файла: {e}")


if __name__ == "__main__":
    print()
    print("==============================")
    print("  MOEX Report Generator")
    print("==============================")
    print()
    print("Выберите режим:")
    print("  1 — Отчёт за прошлую неделю (авто)")
    print("  2 — Указать даты вручную")
    print()

    choice = input("> ").strip()

    if choice == "1":
        start_input, end_input = get_previous_week_dates()
        print(f"Диапазон: {start_input} – {end_input}")
        print()
    elif choice == "2":
        print()
        start_input = input("Стартовая дата (гггг-мм-дд): ").strip()
        end_input = input("Конечная дата (гггг-мм-дд): ").strip()
    else:
        print("Ошибка: выберите 1 или 2")
        input("Нажмите Enter для выхода...")
        exit(1)

    for inp in (start_input, end_input):
        try:
            datetime.strptime(inp, "%Y-%m-%d")
        except ValueError:
            print(f"Ошибка: дата '{inp}' не соответствует формату гггг-мм-дд")
            input("Нажмите Enter для выхода...")
            exit(1)

    if start_input > end_input:
        print("Ошибка: стартовая дата не может быть позже конечной")
        input("Нажмите Enter для выхода...")
        exit(1)

    START_DATE = start_input
    END_DATE = end_input

    start_dt = datetime.strptime(START_DATE, "%Y-%m-%d")
    end_dt = datetime.strptime(END_DATE, "%Y-%m-%d")
    LOAD_FROM = (start_dt - timedelta(days=14)).strftime("%Y-%m-%d")
    LOAD_TO = END_DATE

    main()

    print()
    input("Готово. Нажмите Enter для выхода...")
